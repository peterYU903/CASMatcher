import streamlit as st
import pandas as pd
import numpy as np
import zipfile, pymupdf, re, os
from docx import Document

tree_data = []

def get_standard_names():
    standard_names = os.listdir('standards/')
    return standard_names

def zip_outputs():
    with zipfile.ZipFile('outputs.zip', 'w', zipfile.ZIP_DEFLATED) as zipf:
        for filename in os.listdir('outputs/'):
            file_path = 'outputs/' + filename
            zipf.write(file_path)

def clear_folder():
    files = os.listdir('outputs/')
    for file in files:
        os.remove(os.path.join('outputs/', file))

class CASMatcher:
    def __init__(self):
        return

    def is_single_digit(self, string):
        return len(string) == 1 and string.isdigit()

    def find_numeric_hyphen_strings(self, string):
        pattern_1 = r'^\d+-\d+$'
        pattern_2 = r'^\d+-\d+-\d+$'
        if re.match(pattern_1, string) or re.match(pattern_2, string):
            return True
        return False

    def fit_excel(self, writer):
        for sheet_name in writer.sheets:
            worksheet = writer.sheets[sheet_name]
            for col in worksheet.columns:
                max_length = 0
                col_name = col[0].column_letter
                for cell in col:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                worksheet.column_dimensions[col_name].width = max_length

    def get_MDSReport_docx(self, report):
        doc = Document(report)
        data = []
        for table in doc.tables:
            for row in table.rows:
                temp = row.cells[0].text.strip(" |-,")
                if self.is_single_digit(temp):
                    data.append([cell.text for cell in row.cells])
        data = np.array(data)[:,0:3]
        data = np.char.strip(data.astype(str), " |-,")
        df = pd.DataFrame(data)
        df = df.rename(columns={df.columns[0]: 'Level', df.columns[1]: 'Substance Name', df.columns[2]: 'CAS Number'})
        return df

    def get_MDSReport_pdf(self, report):
        doc = pymupdf.Document(report)
        dfs = []
        for page in doc:
            tabs = page.find_tables()
            if len(tabs.tables):
                for tab in tabs:
                    data = np.array(tab.extract())[:,0:3]
                    data = np.char.strip(data.astype(str), " |-,")
                    df = pd.DataFrame(data)
                    df = df[df.iloc[:, 0].apply(self.is_single_digit)]
                    df = df.rename(columns={df.columns[0]: 'Level', df.columns[1]: 'Substance Name', df.columns[2]: 'CAS Number'})
                    dfs.append(df)
        merged_df = pd.concat(dfs, ignore_index=True)
        return merged_df

    def preprocess_compareList(self, standard_list):
        df = pd.read_excel(standard_list)
        df = df.loc[:, ['CAS Number', 'Chemical Name']]
        df['CAS Number'] = df['CAS Number'].str.strip(" ,")
        df['CAS Number'] = df['CAS Number'].str.split(',')
        df = df.explode('CAS Number').reset_index(drop=True)
        df.to_excel('standards/' + 'processed_' + standard_list.name.split('/')[-1], index=False)
        return df

    def get_result(self, report, standard_name):
        file_type = report.name.split('.')[-1]
        compareList = pd.read_excel('standards/' + standard_name)
        if file_type == 'docx':
            data_MDS = self.get_MDSReport_docx(report)
        else:
            data_MDS = self.get_MDSReport_pdf(report)       
        result = pd.merge(
            data_MDS,
            compareList, on='CAS Number',
            how='left')
        count = result['Chemical Name'].notnull().sum()
        result['Chemical Name'] = result['Chemical Name'].fillna('---')
        level = 0
        yielded_rows = []
        for _, row in result[::-1].iterrows():
            currL = int(row['Level'])
            if row['Chemical Name'] != '---' or level > currL:
                yielded_rows.append(row)
                level = currL
        result = result[result.iloc[:, 2].apply(self.find_numeric_hyphen_strings)]
        result_Found = pd.DataFrame(yielded_rows[::-1])
        filename = report.name.split('/')[-1].split('.')[0] + '.xlsx'
        filepath = 'outputs/' + standard_name.split('_')[-1].split('.')[0] + '_' + str(count) + "_" + filename
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            result.to_excel(writer, sheet_name="Total List", index=False)
            if count:
                result_Found.to_excel(writer, sheet_name="Summary", index=False)
                tree_data.append((standard_name.split('_')[-1].split('.')[0], str(count), filename))
            self.fit_excel(writer)

def main():
    matcher = CASMatcher()
    st.title('CASMatcher Application')
    st.header('1. Upload the standard lists for comparison:')
    standard_lists = st.file_uploader(
        label='Upload Standard Lists',
        type=["xlsx"],
        accept_multiple_files=True,
        )
    if standard_lists is not None:
        for standard_list in standard_lists:
            matcher.preprocess_compareList(standard_list)
    st.header('2. Choose the list for matching:')
    standard_names = get_standard_names()
    standard_name = st.selectbox(
        label='Select Standard List',
        options=standard_names,
        )
    st.header('3. Upload the MAS reports:')
    if standard_name is not None:
        MDSreports = st.file_uploader(
            label='Upload MAS Reports',
            type=["docx", "pdf"],
            accept_multiple_files=True,
            )
        if MDSreports is not None:
            for report in MDSreports:
                matcher.get_result(report, standard_name)
        zip_outputs()
    else:
        st.caption('Please upload and select standard list for further actions.')
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Download the output files:")
        with open('outputs.zip', 'rb') as datazip:
            st.download_button(
                label='Download ZIP',
                data=datazip,
                file_name="outputs.zip",
                mime="application/octet-stream"
                )
    with col2:
        st.subheader("Clear the output files:")
        if st.button(label='Clear Outputs'):
            clear_folder()


if __name__ == '__main__':
    main()