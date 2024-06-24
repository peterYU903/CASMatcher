import streamlit as st
import pandas as pd
import numpy as np
import zipfile, pymupdf, re, os, tempfile
from docx import Document

def get_standard_names():
    standard_names = os.listdir('./standards/')
    return standard_names

def zip_outputs():
    if os.path.exists('./outputs.zip'):
        os.remove('./outputs.zip')
    with zipfile.ZipFile('./outputs.zip', 'w') as zipf:
        for filename in os.listdir('./outputs/'):
            if filename.split('.')[-1] != 'txt':
                file_path = './outputs/' + filename
                zipf.write(file_path, compress_type=zipfile.ZIP_DEFLATED)

class CASMatcher:
    def __init__(self):
        return

    def is_single_digit(self, string: str) -> bool:
        return len(string) == 1 and string.isdigit()

    def find_numeric_hyphen_strings(self, string: str) -> bool:
        pattern_1 = r'^\d+-\d+$'
        pattern_2 = r'^\d+-\d+-\d+$'
        if re.match(pattern_1, string) or re.match(pattern_2, string):
            return True
        return False

    def fit_excel(self, writer: pd.ExcelWriter):
        for sheet_name in writer.sheets:
            worksheet = writer.sheets[sheet_name]
            for col in worksheet.columns:
                max_length = 0
                col_name = col[0].column_letter
                for cell in col:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                worksheet.column_dimensions[col_name].width = max_length

    def get_MDSReport_docx(self, report):
        doc = Document(report)
        data = []
        for table in doc.tables:
            for row in table.rows:
                temp = row.cells[0].text.strip(" |-,")
                if self.is_single_digit(temp):
                    data.append([cell.text for cell in row.cells])
        data = np.array(data)[:,0:3]
        data = np.char.strip(data.astype(str), " |-,")
        df = pd.DataFrame(data)
        df = df.rename(columns={df.columns[0]: 'Level', df.columns[1]: 'Substance Name', df.columns[2]: 'CAS Number'})
        return df

    def get_MDSReport_pdf(self, report):
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file.write(report.read())
        temp_file.close()
        doc = pymupdf.Document(temp_file.name)
        dfs = []
        for page in doc:
            tabs = page.find_tables()
            if len(tabs.tables):
                for tab in tabs:
                    data = np.array(tab.extract())[:,0:3]
                    data = np.char.strip(data.astype(str), " |-,")
                    df = pd.DataFrame(data)
                    df = df[df.iloc[:, 0].apply(self.is_single_digit)]
                    df = df.rename(columns={df.columns[0]: 'Level', df.columns[1]: 'Substance Name', df.columns[2]: 'CAS Number'})
                    dfs.append(df)
        os.remove(temp_file.name)
        merged_df = pd.concat(dfs, ignore_index=True)
        return merged_df

    def preprocess_compareList(self, standard_list):
        df = pd.read_excel(standard_list)
        df = df.loc[:, ['CAS Number', 'Chemical Name']]
        df['CAS Number'] = df['CAS Number'].str.strip(" ,")
        df['CAS Number'] = df['CAS Number'].str.split(',')
        df = df.explode('CAS Number').reset_index(drop=True)
        df.to_excel('./standards/' + 'processed_' + standard_list.name.split('/')[-1], index=False)
        return df

    def get_result(self, report, standard_name: str):
        file_type = report.name.split('.')[-1]
        compareList = pd.read_excel('./standards/' + standard_name)
        if file_type == 'docx':
            data_MDS = self.get_MDSReport_docx(report)
        else:
            data_MDS = self.get_MDSReport_pdf(report)       
        result = pd.merge(
            data_MDS,
            compareList, on='CAS Number',
            how='left')
        count = result['Chemical Name'].notnull().sum()
        result['Chemical Name'] = result['Chemical Name'].fillna('---')
        level = 0
        yielded_rows = []
        for _, row in result[::-1].iterrows():
            currL = int(row['Level'])
            if row['Chemical Name'] != '---' or level > currL:
                yielded_rows.append(row)
                level = currL
        result = result[result.iloc[:, 2].apply(self.find_numeric_hyphen_strings)]
        result_Found = pd.DataFrame(yielded_rows[::-1])
        filename = report.name.split('/')[-1].split('.')[0] + '.xlsx'
        filepath = 'outputs/' + standard_name.split('_')[-1].split('.')[0] + '&' + str(count) + "&" + filename
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            result.to_excel(writer, sheet_name="Total List", index=False)
            if count:
                result_Found.to_excel(writer, sheet_name="Summary", index=False)
            self.fit_excel(writer)

def main():
    matcher = CASMatcher()
    st.title(' 📋 CASMatcher Application')
    st.header('1. Upload the standard lists for comparison:', divider='rainbow')
    standard_lists = st.file_uploader(
        label='Upload Standard Lists',
        type=["xlsx"],
        accept_multiple_files=True,
        )
    if standard_lists is not None:
        for standard_list in standard_lists:
            matcher.preprocess_compareList(standard_list)
        standard_lists = None
    st.header('2. Choose the list for matching:', divider='rainbow')
    standard_names = get_standard_names()
    standard_name = st.selectbox(
        label='Select Standard List',
        options=standard_names,
        )
    if standard_name is not None:
        st.header('3. Upload the MAS reports:', divider='rainbow')
        MDSreports = st.file_uploader(
            label='Upload MAS Reports',
            type=["docx", "pdf"],
            accept_multiple_files=True,
            )
        _, _, body_col3, _, _ = st.columns(5)
        with body_col3:
            process = st.button("Process")
        progress_bar = st.progress(0, text="")
        if process and MDSreports is not None:
            if len(MDSreports) == 0:
                st.warning(" ❌ No reports are uploaded.")
            else:
                progress_bar.progress(0)
                with st.spinner("Operation in progress. Please wait..."):
                    for i, report in enumerate(MDSreports):
                        matcher.get_result(report, standard_name)
                        progress_bar.progress((i + 1) / len(MDSreports))
                st.success(" ✅ Operation Finished. Get the outputs in \"Outputs\" page.")
            MDSreports = None
    else:
        st.caption(' ❗ Please upload and select standard list for further actions.')

if __name__ == '__main__':
    st.set_page_config(page_title='CASMatcher', page_icon='./sources/johnson.jpg')
    st.logo(image='./sources/johnson.jpg')
    main()